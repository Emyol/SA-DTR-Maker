"""Parsing, grouping, and Word-generation logic.

Every function here accepts either a filesystem path or an in-memory
file-like object (BytesIO) wherever openpyxl/python-docx allow it, so the
same code powers both the CLI and the web upload flow without touching disk.
"""

import copy
import datetime as dt
import io
import zipfile
from collections import defaultdict

import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.utils.exceptions import InvalidFileException
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .defaults import (
    ENTRY_FONT_PT,
    ASSIGNED_LABELS,
    WEEKDAY_NAMES,
    TITLE_MARK,
    MAX_LABEL_SEARCH_ROWS,
    MAX_DATA_ROWS,
)


# --------------------------------------------------------------------------- #
#  Errors
# --------------------------------------------------------------------------- #
class DTRError(Exception):
    """Base class for user-facing DTR errors — safe to show verbatim."""


class WorkbookReadError(DTRError):
    """The uploaded file could not be read as a usable .xlsx workbook."""


class NoEntriesError(DTRError):
    """No usable Monday-Saturday time entries were found in the workbook."""


# --------------------------------------------------------------------------- #
#  Parsing the spreadsheet
# --------------------------------------------------------------------------- #
DATE_TEXT_FORMATS = ("%m/%d/%Y", "%m/%d/%y", "%m-%d-%Y", "%m-%d-%y", "%Y-%m-%d")
TIME_TEXT_FORMATS = ("%I:%M %p", "%H:%M", "%I:%M:%S %p", "%H:%M:%S")


def as_time(v):
    if isinstance(v, dt.datetime):
        return v.time()
    if isinstance(v, dt.time):
        return v
    if isinstance(v, str):
        s = v.strip()
        for fmt in TIME_TEXT_FORMATS:
            try:
                return dt.datetime.strptime(s, fmt).time()
            except ValueError:
                continue
    return None


def as_date(v):
    if isinstance(v, dt.datetime):
        return v.date()
    if isinstance(v, dt.date):
        return v
    if isinstance(v, str):
        s = v.strip()
        for fmt in DATE_TEXT_FORMATS:
            try:
                return dt.datetime.strptime(s, fmt).date()
            except ValueError:
                continue
    return None


def find_label_cell(ws, label, max_rows=MAX_LABEL_SEARCH_ROWS):
    """Search near the top of the sheet for a cell whose text matches
    `label` (case-insensitive). Returns (row, col) or None. Bounded to the
    first `max_rows` so a spreadsheet with a huge/forged dimensions range
    can't force an unbounded scan."""
    target = label.strip().upper()
    scan_limit = min(ws.max_row or 1, max_rows)
    for row in ws.iter_rows(max_row=scan_limit):
        for cell in row:
            if cell.value is not None and str(cell.value).strip().upper() == target:
                return cell.row, cell.column
    return None


def detect_layout(ws):
    """Find the ONLINE/ONSITE block columns, header row, and name cell by
    searching for the "ONLINE"/"ONSITE" labels rather than assuming a fixed
    B2/B3/row-5 layout. Tolerates sheets where everything is shifted by a
    row or column. Falls back to the original fixed layout if the labels
    can't be found."""
    blocks = []          # (date_col, in_col, out_col, source, data_start_row)
    name = ""
    for label, source in (("ONLINE", "ONLINE"), ("ONSITE", "ONSITE")):
        loc = find_label_cell(ws, label)
        if loc is None:
            continue
        label_row, col = loc
        date_col = get_column_letter(col)
        in_col = get_column_letter(col + 1)
        out_col = get_column_letter(col + 2)
        header_row = label_row + 1
        data_row = header_row + 1
        blocks.append((date_col, in_col, out_col, source, data_row))
        if not name:
            name_val = ws.cell(row=label_row - 1, column=col).value
            if name_val:
                name = str(name_val).strip()

    if not blocks:
        blocks = [("B", "C", "D", "ONLINE", 5), ("F", "G", "H", "ONSITE", 5)]
    if not name:
        name = str(ws["B2"].value or "").strip()
    return name, blocks


def parse_workbook(source):
    """Parse a spreadsheet into (name, entries, anomalies).

    `source` may be a filesystem path (str) or a file-like object (e.g.
    io.BytesIO) — anything openpyxl.load_workbook accepts.
    """
    try:
        wb = openpyxl.load_workbook(source, data_only=True)
    except zipfile.BadZipFile as exc:
        raise WorkbookReadError(
            "That file isn't a valid .xlsx workbook. It may be corrupted, "
            "saved in an older format (.xls), or password-protected — "
            "try re-saving it as a plain .xlsx without a password."
        ) from exc
    except InvalidFileException as exc:
        raise WorkbookReadError(f"Unsupported file: {exc}") from exc
    except KeyError as exc:
        raise WorkbookReadError(
            "The workbook is missing required internal parts and could not be read."
        ) from exc

    ws = wb.active
    if ws is None or not ws.max_row:
        raise WorkbookReadError("The workbook's first sheet appears to be empty.")

    name, blocks = detect_layout(ws)
    data_start = min(b[4] for b in blocks)

    entries = []          # list of dicts
    anomalies = []        # notes for review

    end_row = min(ws.max_row, data_start + MAX_DATA_ROWS)
    truncated = ws.max_row > end_row

    for row in range(data_start, end_row + 1):
        for date_col, in_col, out_col, source, block_start in blocks:
            if row < block_start:
                continue
            d = as_date(ws[f"{date_col}{row}"].value)
            tin = as_time(ws[f"{in_col}{row}"].value)
            tout = as_time(ws[f"{out_col}{row}"].value)
            if d is None:
                continue
            if tin is None or tout is None:
                anomalies.append(
                    f"Row {row} {source}: date {d} has a missing "
                    f"time-in/time-out (in={tin}, out={tout}) - skipped."
                )
                continue
            start = dt.datetime.combine(d, tin)
            end = dt.datetime.combine(d, tout)
            hours = (end - start).total_seconds() / 3600.0
            if hours <= 0:
                anomalies.append(
                    f"Row {row} {source}: {d} out<=in (in={tin}, out={tout})."
                )
            entries.append({
                "date": d,
                "weekday": d.weekday(),          # Mon=0 .. Sun=6
                "time_in": tin,
                "time_out": tout,
                "hours": hours,
                "source": source,
                "assigned": ASSIGNED_LABELS.get(source, source),
            })

    if truncated:
        anomalies.append(
            f"Workbook has more than {MAX_DATA_ROWS} data rows; only the "
            f"first {MAX_DATA_ROWS} (rows {data_start}-{end_row}) were processed."
        )

    return name, entries, anomalies


# --------------------------------------------------------------------------- #
#  Grouping into weeks (Mon-Sat)
# --------------------------------------------------------------------------- #
def group_into_weeks(entries):
    weeks = defaultdict(list)   # monday_date -> entries
    for e in entries:
        monday = e["date"] - dt.timedelta(days=e["weekday"])
        weeks[monday].append(e)

    result = []
    for monday in sorted(weeks):
        saturday = monday + dt.timedelta(days=5)
        days = defaultdict(list)   # weekday index -> entries
        for e in weeks[monday]:
            days[e["weekday"]].append(e)
        for wd in days:
            days[wd].sort(key=lambda x: x["time_in"])
        total = sum(e["hours"] for e in weeks[monday])
        result.append({
            "monday": monday,
            "saturday": saturday,
            "days": days,
            "total": total,
        })
    return result


# --------------------------------------------------------------------------- #
#  Formatting helpers
# --------------------------------------------------------------------------- #
def fmt_time(t):
    s = t.strftime("%I:%M %p")
    return s[1:] if s.startswith("0") else s


def fmt_date(d):
    return d.strftime("%m/%d/%y")


def fmt_hours(h):
    h = round(h, 2)
    if abs(h - round(h)) < 1e-9:
        return str(int(round(h)))
    return ("%.2f" % h).rstrip("0").rstrip(".")


# --------------------------------------------------------------------------- #
#  Preview (text) output - lets us verify the logic without opening Word
# --------------------------------------------------------------------------- #
def build_preview_text(name, weeks, anomalies):
    lines = []
    lines.append("DTR PREVIEW")
    lines.append("=" * 60)
    lines.append(f"Name: {name}")
    lines.append(f"Weeks with records: {len(weeks)}")
    lines.append("")
    for w in weeks:
        lines.append("-" * 60)
        lines.append(
            f"WEEK: {w['monday']:%b %d, %Y} (Mon) -> "
            f"{w['saturday']:%b %d, %Y} (Sat)   "
            f"Weekly total: {fmt_hours(w['total'])} hrs"
        )
        for wd in range(6):   # Mon..Sat
            day_entries = w["days"].get(wd, [])
            if not day_entries:
                continue
            d = day_entries[0]["date"]
            lines.append(f"   {WEEKDAY_NAMES[wd]:<9} {fmt_date(d)}")
            for e in day_entries:
                lines.append(
                    f"       {fmt_time(e['time_in'])} - {fmt_time(e['time_out'])}"
                    f"   {fmt_hours(e['hours'])} hrs   [{e['source']}] "
                    f"{e['assigned']}"
                )
        # Sunday check
        if w["days"].get(6):
            lines.append("   !! Sunday entries found (no Sunday row on the form):")
            for e in w["days"][6]:
                lines.append(f"       {fmt_date(e['date'])} {fmt_time(e['time_in'])}"
                             f"-{fmt_time(e['time_out'])} [{e['source']}]")
    if anomalies:
        lines.append("")
        lines.append("ANOMALIES / NOTES")
        lines.append("=" * 60)
        for a in anomalies:
            lines.append(" - " + a)

    return "\n".join(lines)


def write_preview(name, weeks, anomalies, path):
    text = build_preview_text(name, weeks, anomalies)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return text


# --------------------------------------------------------------------------- #
#  Word generation
# --------------------------------------------------------------------------- #
def ptext(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t")))


def set_paragraph_text(p_el, text, size_pt=None, bold=None):
    """Replace a paragraph's text, keeping the first run's formatting.
    Optionally override font size (points) and bold."""
    runs = p_el.findall(qn("w:r"))
    rpr = None
    if runs:
        found = runs[0].find(qn("w:rPr"))
        if found is not None:
            rpr = copy.deepcopy(found)
    for r in runs:
        p_el.remove(r)
    new_r = OxmlElement("w:r")
    if rpr is None:
        rpr = OxmlElement("w:rPr")
    if bold:
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))
    if size_pt is not None:
        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), str(int(size_pt * 2)))
    new_r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    p_el.append(new_r)


def set_cell_lines(cell, lines, size_pt=ENTRY_FONT_PT):
    """Write one paragraph per line into a table cell, clearing existing text
    while preserving the cell's original paragraph alignment and run style."""
    orig_first = cell.paragraphs[0]
    ppr = orig_first._element.find(qn("w:pPr"))
    ppr = copy.deepcopy(ppr) if ppr is not None else None
    rpr = None
    if orig_first.runs:
        found = orig_first.runs[0]._element.find(qn("w:rPr"))
        if found is not None:
            rpr = copy.deepcopy(found)

    # remove extra paragraphs, keep the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)

    def ensure_ppr(paragraph):
        if ppr is not None and paragraph._element.find(qn("w:pPr")) is None:
            paragraph._element.insert(0, copy.deepcopy(ppr))

    def add_line(paragraph, text):
        ensure_ppr(paragraph)
        run = paragraph.add_run(str(text))
        if rpr is not None:
            run._element.insert(0, copy.deepcopy(rpr))
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not lines:
        return
    add_line(first, lines[0])
    for extra in lines[1:]:
        p = cell.add_paragraph()
        add_line(p, extra)


def make_page_break():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _append_run(p_el, text, base_rpr=None, size_pt=None, bold=None):
    new_r = OxmlElement("w:r")
    rpr = copy.deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")
    if bold and rpr.find(qn("w:b")) is None:
        rpr.append(OxmlElement("w:b"))
    if size_pt is not None:
        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), str(int(size_pt * 2)))
    new_r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    p_el.append(new_r)


def fill_total_box(block_els, total_str):
    """Append the weekly total to the floating 'Total Hours Rendered:' box(es),
    keeping the original label text/format and styling only the value (12pt bold).
    Word stores a modern + fallback copy, so we update every matching box."""
    for el in block_els:
        for txbx in el.iter(qn("w:txbxContent")):
            for p in txbx.findall(qn("w:p")):
                if "Total Hours Rendered" in ptext(p):
                    runs = p.findall(qn("w:r"))
                    base = runs[0].find(qn("w:rPr")) if runs else None
                    _append_run(p, f" {total_str} Hours",
                                base_rpr=base, size_pt=12, bold=True)


def find_week_grid(block_els, doc):
    tables = [el for el in block_els if el.tag == qn("w:tbl")]
    if not tables:
        return None
    return Table(tables[0], doc)


def fill_week_block(block_els, doc, ident, week):
    name_line = (f"Name  {ident['name']}          Term  {ident['term']}"
                 f"          School Year  {ident['school_year']}")
    office_line = (f"Office  {ident['office']}          "
                   f"Supervisor  {ident['supervisor']}")
    month_line = (f"Month  {week['monday']:%B}   from  {week['monday']:%m/%d}"
                  f"   to  {week['saturday']:%m/%d}   Year  {week['monday']:%Y}")

    for el in block_els:
        if el.tag == qn("w:p"):
            txt = ptext(el)
            if "School Year" in txt:
                set_paragraph_text(el, name_line)
            elif "Office" in txt and "Supervisor" in txt:
                set_paragraph_text(el, office_line)
            elif txt.strip().startswith("Month"):
                set_paragraph_text(el, month_line)

    fill_total_box(block_els, fmt_hours(week["total"]))

    grid = find_week_grid(block_els, doc)
    if grid is None:
        return

    # rows 1..6 correspond to Monday..Saturday
    for wd in range(6):
        day_entries = week["days"].get(wd, [])
        if not day_entries:
            continue
        row = grid.rows[wd + 1]
        cells = row.cells
        d = day_entries[0]["date"]
        set_cell_lines(cells[0], [WEEKDAY_NAMES[wd], fmt_date(d)])
        set_cell_lines(cells[1], [fmt_time(e["time_in"]) for e in day_entries])
        set_cell_lines(cells[2], [fmt_time(e["time_out"]) for e in day_entries])
        set_cell_lines(cells[3], [fmt_hours(e["hours"]) for e in day_entries])
        set_cell_lines(cells[4], [e["assigned"] for e in day_entries])

    # figure out which column is "# OF HOURS RENDERED" from the header row
    hours_col = 3
    for i, c in enumerate(grid.rows[0].cells):
        if "HOURS RENDERED" in c.text.upper():
            hours_col = i
            break

    # weekly total: find the row whose first cell says "Hours Rendered"
    for row in grid.rows:
        cells = row.cells
        first_txt = cells[0].text.strip().lower()
        if first_txt.startswith("hours rendered"):
            target = cells[hours_col] if hours_col < len(cells) else cells[1]
            set_cell_lines(target, [fmt_hours(week["total"])])
            break


def split_two_blocks(unit_els):
    """Split a two-form page unit into [top_block_els, bottom_block_els]."""
    idxs = [i for i, el in enumerate(unit_els)
            if el.tag == qn("w:p") and TITLE_MARK in ptext(el)]
    if len(idxs) >= 2:
        return [unit_els[:idxs[1]], unit_els[idxs[1]:]]
    return [unit_els]


def strip_trailing_empties(els):
    """Drop trailing blank paragraphs so an added page break can't overflow."""
    while els and els[-1].tag == qn("w:p") and ptext(els[-1]).strip() == "":
        els.pop()
    return els


def insert_before(body, el, sectPr):
    if sectPr is not None:
        sectPr.addprevious(el)
    else:
        body.append(el)


def generate_docx(template, ident, weeks, out=None):
    """Fill the DTR template for every week and either save it to `out`
    (path or file-like) or, if `out` is None, return the finished .docx as
    bytes.

    `template` may be a filesystem path or a file-like object.
    `ident` must have name/term/school_year/office/supervisor keys.
    """
    if not weeks:
        raise NoEntriesError(
            "No weekly records to generate — there's nothing to fill in the DTR."
        )

    doc = Document(template)
    body = doc.element.body
    children = list(body)

    sectPr = None
    if children and children[-1].tag == qn("w:sectPr"):
        sectPr = children[-1]
    unit_end = children.index(sectPr) if sectPr is not None else len(children)

    # The whole two-form layout is a proven single page -> use it as the unit.
    unit_template = [copy.deepcopy(el) for el in children[:unit_end]]

    # clear body except the section properties
    for c in list(body):
        if c is not sectPr:
            body.remove(c)

    num_pages = (len(weeks) + 1) // 2
    for page_idx in range(num_pages):
        unit_copy = [copy.deepcopy(el) for el in unit_template]
        blocks = split_two_blocks(unit_copy)

        week_a = weeks[page_idx * 2]
        fill_week_block(blocks[0], doc, ident, week_a)

        b_index = page_idx * 2 + 1
        if b_index < len(weeks) and len(blocks) > 1:
            fill_week_block(blocks[1], doc, ident, weeks[b_index])
        # if there's no second week, the bottom form stays as a blank template

        unit_copy = strip_trailing_empties(unit_copy)
        for el in unit_copy:
            insert_before(body, el, sectPr)
        if page_idx != num_pages - 1:
            insert_before(body, make_page_break(), sectPr)

    if out is not None:
        doc.save(out)
        return None

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
